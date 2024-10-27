import gradio as gr
from pdf2docx import Converter
from docx import Document
import os
import glob
import base64
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xml.etree.ElementTree as ET

def find_ttf_fonts():
    files = glob.glob('**/*.ttf', recursive=True)
    return files

def embed_font_in_html(font_path, font_name, html_content):
    with open(font_path, "rb") as font_file:
        font_data = font_file.read()
    encoded_font = base64.b64encode(font_data).decode('utf-8')
    
    font_style = f"""
    <style>
    @font-face {{
        font-family: '{font_name}';
        src: url(data:font/ttf;base64,{encoded_font}) format('truetype');
    }}
    body {{
        font-family: '{font_name}', Arial, sans-serif;
        margin: 0;
        padding: 0;
        background-color: white;
    }}
    .page {{
        position: relative;
        width: 8.5in;
        margin: 20px auto;
        padding: 20px;
        box-sizing: border-box;
        background-color: white;
        box-shadow: 0 0 10px rgba(0,0,0,0.1);
    }}
    .paragraph {{
        margin: 0;
        padding: 0;
        position: relative;
    }}
    .image-container {{
        display: inline-block;
        position: relative;
        vertical-align: middle;
    }}
    img {{
        max-width: 100%;
        height: auto;
        display: inline-block;
        vertical-align: middle;
    }}
    table {{
        border-collapse: collapse;
        width: 100%;
        margin: 10px 0;
    }}
    td, th {{
        border: 1px solid black;
        padding: 8px;
        position: relative;
    }}
    </style>
    """
    return font_style + html_content

def extract_images_from_doc(doc):
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                image_type = rel.target_part.content_type.split('/')[-1]
                if image_type.lower() not in ['jpeg', 'jpg', 'png', 'gif']:
                    image_type = 'png'
                encoded_image = base64.b64encode(image_data).decode('utf-8')
                images[rel.rId] = f"data:image/{image_type};base64,{encoded_image}"
            except Exception as e:
                print(f"Error processing image: {str(e)}")
                continue
    return images

def get_image_position(element):
    try:
        anchor = element.find('.//wp:anchor', 
            {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if anchor is not None:
            pos_h = anchor.find('.//wp:positionH', 
                {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            pos_v = anchor.find('.//wp:positionV', 
                {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            
            if pos_h is not None and pos_v is not None:
                x = pos_h.find('.//wp:posOffset', 
                    {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                y = pos_v.find('.//wp:posOffset', 
                    {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                
                if x is not None and y is not None:
                    return {
                        'x': int(x.text) / 914400,
                        'y': int(y.text) / 914400
                    }
    except Exception:
        pass
    return None

def process_paragraph(paragraph, images_dict):
    html_content = '<div class="paragraph">'
    
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        html_content += '<div style="text-align: center;">'
    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        html_content += '<div style="text-align: right;">'
    else:
        html_content += '<div>'
    
    for run in paragraph.runs:
        style = []
        if run.bold: style.append('font-weight: bold')
        if run.italic: style.append('font-style: italic')
        if run.underline: style.append('text-decoration: underline')
        if run.font.size: style.append(f'font-size: {run.font.size.pt}pt')
        
        drawing_elements = run._element.findall('.//w:drawing',
            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        for drawing in drawing_elements:
            blip = drawing.find('.//a:blip',
                {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            if blip is not None:
                image_rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if image_rel_id in images_dict:
                    position = get_image_position(drawing)
                    if position:
                        style_pos = f"position: absolute; left: {position['x']}in; top: {position['y']}in;"
                        html_content += f'<div class="image-container" style="{style_pos}">'
                        html_content += f'<img src="{images_dict[image_rel_id]}" alt="Document Image"/>'
                        html_content += '</div>'
                    else:
                        html_content += f'<div class="image-container">'
                        html_content += f'<img src="{images_dict[image_rel_id]}" alt="Document Image"/>'
                        html_content += '</div>'
        
        style_str = '; '.join(style)
        if run.text.strip():
            html_content += f'<span style="{style_str}">{run.text}</span>'
    
    html_content += '</div></div>'
    return html_content

def process_table(table, images_dict):
    html_content = '<table>'
    for row in table.rows:
        html_content += '<tr>'
        for cell in row.cells:
            html_content += '<td>'
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    style = []
                    if run.bold: style.append('font-weight: bold')
                    if run.italic: style.append('font-style: italic')
                    if run.underline: style.append('text-decoration: underline')
                    if run.font.size: style.append(f'font-size: {run.font.size.pt}pt')
                    
                    drawing_elements = run._element.findall('.//w:drawing',
                        {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    
                    for drawing in drawing_elements:
                        blip = drawing.find('.//a:blip',
                            {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                        if blip is not None:
                            image_rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if image_rel_id in images_dict:
                                html_content += f'<div class="image-container">'
                                html_content += f'<img src="{images_dict[image_rel_id]}" alt="Table Cell Image"/>'
                                html_content += '</div>'
                    
                    style_str = '; '.join(style)
                    if run.text.strip():
                        html_content += f'<span style="{style_str}">{run.text}</span>'
            html_content += '</td>'
        html_content += '</tr>'
    html_content += '</table>'
    return html_content

def pdf_to_html(pdf_file, font_name):
    if not pdf_file:
        return None
        
    try:
        docx_filename = pdf_file.name.replace('.pdf', '.docx')
        cv = Converter(pdf_file.name)
        cv.convert(docx_filename)
        cv.close()
        
        doc = Document(docx_filename)
        images_dict = extract_images_from_doc(doc)
        
        html_content = """<!DOCTYPE html>
        <html>
        <head>
            <meta charset='utf-8'>
            <title>Converted Document</title>
        </head>
        <body>
        <div class="page">
        """
        
        paragraph_map = {}
        current_paragraph_index = 0
        for para in doc.paragraphs:
            paragraph_map[para._element] = current_paragraph_index
            current_paragraph_index += 1
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                if element in paragraph_map:
                    paragraph = doc.paragraphs[paragraph_map[element]]
                    html_content += process_paragraph(paragraph, images_dict)
            elif element.tag.endswith('tbl'):
                table_index = len([e for e in doc.element.body[:doc.element.body.index(element)] 
                                 if e.tag.endswith('tbl')])
                html_content += process_table(doc.tables[table_index], images_dict)
        
        html_content += "</div></body></html>"
        
        ttf_files = {os.path.basename(f): f for f in find_ttf_fonts()}
        if font_name in ttf_files:
            font_path = ttf_files[font_name]
            font_name_clean = os.path.splitext(font_name)[0]
            html_content = embed_font_in_html(font_path, font_name_clean, html_content)
        
        html_filename = "output_with_font.html"
        with open(html_filename, "w", encoding="utf-8") as html_file:
            html_file.write(html_content)
        
        os.remove(docx_filename)
        return html_filename
        
    except Exception as e:
        print(f"Error in pdf_to_html: {str(e)}")
        return None

# Gradio Interface
with gr.Blocks(theme=gr.themes.Soft()) as app:
    gr.Markdown("# Bionic Reading PDF Converter")
    
    with gr.Row():
        gr.Image("image.jpeg", 
                label="Bionic Reading Example", 
                show_label=False,
                width=400,
                height=300)


    with gr.Row():
        with gr.Column(scale=2):
            pdf_input = gr.File(
                label="Upload Your PDF",
                file_types=[".pdf"],
                file_count="single"
            )
            
            ttf_files = find_ttf_fonts()
            font_dropdown = gr.Dropdown(
                [os.path.basename(font) for font in ttf_files],
                label="Select Font Style",
                value=os.path.basename(ttf_files[0]) if ttf_files else None,
                info="Choose your preferred reading font"
            )
            
            convert_pdf_to_html = gr.Button(
                "Convert to Bionic Format",
                variant="primary",
                size="lg"
            )
            
            font_output = gr.File(
                label="Download Enhanced HTML File",
                type="filepath"
            )

    with gr.Row():
        example_files = [
            os.path.join("examples", f) 
            for f in os.listdir("examples") 
            if f.endswith('.pdf')
        ] if os.path.exists("examples") else []
        
        if example_files:
            gr.Examples(
                example_files,
                pdf_input,
                label="Sample PDFs"
            )
    
    with gr.Row():
        gr.Markdown(
            """
            ---
            📝 Best results with text-based PDFs (not scanned documents)
            """
        )
            
    convert_pdf_to_html.click(
        pdf_to_html,
        inputs=[pdf_input, font_dropdown],
        outputs=[font_output]
    )

app.launch(debug=True)
